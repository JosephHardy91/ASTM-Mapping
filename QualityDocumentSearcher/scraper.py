# Title Scraper
import os, re, csv
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.converter import XMLConverter, HTMLConverter, TextConverter
from pdfminer.layout import LAParams
from cStringIO import StringIO
from collections import defaultdict
import docx
from searcher import Searcher
from compoundfiles import CompoundFileReader, CompoundFileError
from struct import unpack
import xlrd
import openpyxl
import jellyfish as jf


def doc2text(path):
    text = u""
    cr = CompoundFileReader(path)
    # Load WordDocument stream:
    try:
        f = cr.open("WordDocument")
        doc = f.read()
        f.close()
    except:
        cr.close()
        raise CompoundFileError, "The file is corrupted or it is not a Word document at all."
    # Extract file information block and piece table stream informations from it:
    fib = doc[:1472]
    fcClx = unpack("L", fib[0x01a2l:0x01a6l])[0]
    lcbClx = unpack("L", fib[0x01a6l:0x01a6 + 4l])[0]
    tableFlag = unpack("L", fib[0x000al:0x000al + 4l])[0] & 0x0200l == 0x0200l
    tableName = ("0Table", "1Table")[tableFlag]
    # Load piece table stream:
    try:
        f = cr.open(tableName)
        table = f.read()
        f.close()
    except:
        cr.close();
        raise CompoundFileError, "The file is corrupt. '%s' piece table stream is missing." % tableName
    cr.close()
    # Find piece table inside a table stream:
    clx = table[fcClx:fcClx + lcbClx]
    pos = 0
    pieceTable = ""
    lcbPieceTable = 0
    while True:
        if clx[pos] == "\x02":
            # This is piece table, we store it:
            lcbPieceTable = unpack("l", clx[pos + 1:pos + 5])[0]
            pieceTable = clx[pos + 5:pos + 5 + lcbPieceTable]
            break
        elif clx[pos] == "\x01":
            # This is beggining of some other substructure, we skip it:
            pos = pos + 1 + 1 + ord(clx[pos + 1])
        else:
            break
    if not pieceTable: raise CompoundFileError, "The file is corrupt. Cannot locate a piece table."
    # Read info from pieceTable, about each piece and extract it from WordDocument stream:
    pieceCount = (lcbPieceTable - 4) / 12
    for x in xrange(pieceCount):
        cpStart = unpack("l", pieceTable[x * 4:x * 4 + 4])[0]
        cpEnd = unpack("l", pieceTable[(x + 1) * 4:(x + 1) * 4 + 4])[0]
        ofsetDescriptor = ((pieceCount + 1) * 4) + (x * 8)
        pieceDescriptor = pieceTable[ofsetDescriptor:ofsetDescriptor + 8]
        fcValue = unpack("L", pieceDescriptor[2:6])[0]
        isANSII = (fcValue & 0x40000000) == 0x40000000
        fc = fcValue & 0xbfffffff
        cb = cpEnd - cpStart
        enc = ("utf-16", "cp1252")[isANSII]
        cb = (cb * 2, cb)[isANSII]
        text += doc[fc:fc + cb].decode(enc, "ignore")
    return "\n".join(text.splitlines())


class Scraper(object):
    def __init__(self, qdpath, option):
        self.qd_path = qdpath
        self.option = option

    def walk(self):
        self.qd_dict = defaultdict(list)
        osobjs = os.walk(self.qd_path)
        for sub, dir, files in osobjs:
            for file in files:
                if not self.is_txt(file) and self.option == "2b":
                    file = self.convert_to_txt(os.path.join(sub, file))
                    self.qd_dict[sub].append(file)
                elif self.is_txt(file):
                    self.qd_dict[sub].append(file)
                    # if self.is_astm_standard(osobj):
                    #     id, title = re.findall('([D,F]\d+-\d+) (.+)', os.path.splitext(osobj)[0])[0]
                    #     self.titles_dict[id] = title

    def scrape(self):
        self.scrape_results = defaultdict(lambda: defaultdict(list))
        for category_of_document_path in self.qd_dict:
            for file in self.qd_dict[category_of_document_path]:
                s = Searcher(category_of_document_path, file, self.qd_dict)
                s.search()
                self.scrape_results[os.path.split(category_of_document_path)[1]][os.path.splitext(file)[0]] = s.results
        print len(self.scrape_results)

    def organize_results(self):
        self.structured_references = defaultdict(lambda: defaultdict(list))
        self.tagContradictions = []
        self.notReferenced = []
        #print self.scrape_results
        for category in self.scrape_results:
            for file_name in self.scrape_results[category]:
                for tag in self.scrape_results[category][file_name]:
                    finds = self.scrape_results[category][file_name][tag][0]
                    known_files = []
                    if len(self.scrape_results[category][file_name][tag]) > 1:
                        known_files = self.scrape_results[category][file_name][tag][1:]
                    referenced = False
                    for find in finds:
                        idtag, title = find[0], find[1]
                        print idtag, title
                        for tag2 in self.structured_references:
                            for id in self.structured_references[tag2]:
                                if id in idtag or idtag in id:
                                    referenced = True
                                    if len(self.structured_references[tag2][id]) == 0:
                                        if title.strip() != '':
                                            self.structured_references[tag2][id] = [title, {title: [file_name]}]
                                    elif title in self.structured_references[tag2][id][1] or max(
                                            [jf.jaro_winkler(unicode(title), unicode(t)) for t in
                                             self.structured_references[tag2][id][1]]) > 0.85:
                                        if title.strip() != '':
                                            self.structured_references[tag2][id][1][title].append(file_name)
                                    else:
                                        if title.strip() != '':
                                            self.structured_references[tag2][id][1][title] = [file_name]
                                else:
                                    if len(self.structured_references[tag2][id]) != 0:
                                        if title in self.structured_references[tag2][id][1] or max(
                                                [jf.jaro_winkler(unicode(title), unicode(t)) for t in
                                                 self.structured_references[tag2][id][1]]) > 0.85:
                                            if idtag.strip() == '' and title.strip() != '':
                                                referenced = True
                                                self.structured_references[tag2][id][1][title].append(file_name)
                                            elif title.strip() != '':
                                                self.tagContradictions.append(
                                                    [file_name, self.scrape_results[category][file_name], tag2, id])
                    if not referenced:
                        self.notReferenced.append([file_name, self.scrape_results[category][file_name]])
        print len(self.structured_references)
        self.vote_and_restructure()
        self.flip_references()

    def vote_and_restructure(self):
        self.shunned_titles = defaultdict(lambda: defaultdict(list))
        for tag in self.structured_references:
            for id in self.structured_references[tag]:
                votes_dict = {title: len(self.structured_references[tag][id][1][title]) for title in
                              self.structured_references[tag][id][1]}
                votes_list = [(title, votes_dict[title]) for title in
                              sorted(votes_dict, key=votes_dict.get, reverse=True)]
                self.structured_references[tag][id] = [votes_list[0],
                                                       [title for title in self.structured_references[tag][id][1]]]
                self.shunned_titles[tag][id] = votes_list[1:]

    def flip_references(self):
        self.linkage_dictionary = defaultdict(list)
        for tag in self.structured_references:
            for id in self.structured_references[tag]:
                for title in self.structured_references[tag][id][1]:
                    self.linkage_dictionary[title].append(' '.join([tag, id, self.structured_references[tag][id][0]]))

    def output_results(self, output_path):
        with open(output_path, 'wb') as outcsv:
            csvwriter = csv.writer(outcsv, delimiter=',', quoting=csv.QUOTE_MINIMAL, quotechar='"')
            csvwriter.writerow(['Document', 'Links'])
            for title in self.linkage_dictionary:
                csvwriter.writerrow([title, ','.join(self.linkage_dictionary[title])])
                # for category in self.scrape_results:
                #     for file in self.scrape_results[category]:
                #         csvwriter.writerow([category, file, ",".join()])
                # [' '.join([tag, t]) for tag in self.scrape_results[category][file] for c in
                #  self.scrape_results[category][file][tag] for t in c])])

    # def is_astm_reference(self, file_name):
    #     try:
    #         return os.path.splitext(file_name)[1].lower() == '.pdf' \
    #                and len(re.findall('([D,F]\d+-\d+) (.+)', os.path.splitext(file_name)[0])[0]) == 2
    #     except:
    #         return False

    def is_txt(self, file_name):
        return os.path.splitext(file_name.lower())[1] == '.txt'

    def convert_to_txt(self, data):
        extension = os.path.splitext(data)[1].lower()
        if extension == '.pdf':
            fp = file(data, 'rb')
            fname = data
            rsrcmgr = PDFResourceManager()
            retstr = StringIO()
            codec = 'utf-8'
            laparams = LAParams()
            device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
            # Create a PDF interpreter object.
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            # Process each page contained in the document.

            for page in PDFPage.get_pages(fp):
                interpreter.process_page(page)
                data = retstr.getvalue()

            with open(os.path.splitext(fname)[0] + '.txt', 'wb') as output_file:
                output_file.write(data)

            return os.path.splitext(os.path.split(fname)[1])[0] + '.txt'

        elif extension == '.docx':
            document = docx.Document(data)
            text = [paragraph.text for paragraph in document.paragraphs]
            text.extend([cell.text for table in document.tables for row in table.rows for cell in row.cells])
            with open(os.path.splitext(data)[0] + '.txt', 'wb') as output_file:
                output_file.write('\n'.join([c for x in text for c in x if 32 <= ord(c) <= 127]))

            return os.path.splitext(data)[0] + '.txt'

        elif extension == '.doc':
            with open(os.path.splitext(data)[0] + '.txt', 'wb') as output_file:
                output_file.write("\n".join([x for x in doc2text(data) if 32 <= ord(x) <= 127]))

            return os.path.splitext(data)[0] + '.txt'

        elif extension == '.xlsx':
            wb = openpyxl.Workbook(data)
            text = [cell.value for ws in wb for cell in ws.rows]
            with open(os.path.splitext(data)[0] + '.txt', 'wb') as output_file:
                output_file.write("\n".join([x for x in text if 32 <= ord(x) <= 127]))

            return os.path.splitext(data)[0] + '.txt'

        elif extension == '.xls':
            wb = xlrd.open_workbook(data)
            text = [cell.value for ws in wb.sheets() for row in ws.get_rows() for cell in row]
            with open(os.path.splitext(data)[0] + '.txt', 'wb') as output_file:
                output_file.write("\n".join([str(x) for x in text]))

            return os.path.splitext(data)[0] + '.txt'
